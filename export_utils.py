from io import BytesIO
from docx import Document


def _safe_str(value, fallback: str = "") -> str:
    """Convert any value to a plain string, never raising."""
    try:
        return str(value) if value is not None else fallback
    except Exception:
        return fallback


def build_report_docx_bytes(result: dict) -> bytes:
    doc = Document()

    summary = result.get("summary") or {}
    info = result.get("info") or {}
    story_clarity = result.get("story_clarity") or {}
    visual = result.get("visual") or {}
    audio = result.get("audio") or {}
    rows = result.get("rows") or []
    transcript = result.get("transcript") or ""

    # Title
    doc.add_heading("AI Amendment Bot — QC Report", 0)

    # Scores
    doc.add_heading("Scores", level=1)
    doc.add_paragraph(f"Story Clarity: {story_clarity.get('score', 3)}/5")
    doc.add_paragraph(f"Information Clarity: {info.get('score', 3)}/5")
    doc.add_paragraph(f"Visuals: {visual.get('score', 3)}/5")
    doc.add_paragraph(f"Audio: {audio.get('score', 3)}/5")
    doc.add_paragraph(f"Predicted Retention: {summary.get('retention', 'Medium')}")

    # Overall Review
    doc.add_heading("Overall Review", level=1)
    doc.add_paragraph(summary.get("overall_review", ""))

    # Suggestions
    suggestions = summary.get("suggestions", [])
    if suggestions:
        doc.add_heading("Top Suggestions", level=1)
        for s in suggestions:
            doc.add_paragraph(s, style="List Bullet")

    # Information Clarity
    doc.add_heading("Information Clarity", level=1)
    doc.add_paragraph(info.get("summary", ""))

    if info.get("strengths"):
        doc.add_paragraph("Strengths:", style="List Bullet")
        for s in info["strengths"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if info.get("improvements"):
        doc.add_paragraph("Improvements:", style="List Bullet")
        for s in info["improvements"]:
            doc.add_paragraph(s, style="List Bullet 2")

    # Story Clarity
    doc.add_heading("Story Clarity", level=1)
    doc.add_paragraph(story_clarity.get("summary", ""))

    if story_clarity.get("strengths"):
        doc.add_paragraph("Strengths:", style="List Bullet")
        for s in story_clarity["strengths"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if story_clarity.get("improvements"):
        doc.add_paragraph("Improvements:", style="List Bullet")
        for s in story_clarity["improvements"]:
            doc.add_paragraph(s, style="List Bullet 2")

    # Visual Review
    doc.add_heading("Visual Review", level=1)
    doc.add_paragraph(visual.get("summary", ""))

    if visual.get("strengths"):
        doc.add_paragraph("Strengths:", style="List Bullet")
        for s in visual["strengths"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if visual.get("issues"):
        doc.add_paragraph("Issues:", style="List Bullet")
        for s in visual["issues"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if visual.get("suggestions"):
        doc.add_paragraph("Suggestions:", style="List Bullet")
        for s in visual["suggestions"]:
            doc.add_paragraph(s, style="List Bullet 2")

    # Audio Review
    doc.add_heading("Audio Review", level=1)
    doc.add_paragraph(audio.get("summary", ""))

    if audio.get("strengths"):
        doc.add_paragraph("Strengths:", style="List Bullet")
        for s in audio["strengths"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if audio.get("issues"):
        doc.add_paragraph("Issues:", style="List Bullet")
        for s in audio["issues"]:
            doc.add_paragraph(s, style="List Bullet 2")

    if audio.get("suggestions"):
        doc.add_paragraph("Suggestions:", style="List Bullet")
        for s in audio["suggestions"]:
            doc.add_paragraph(s, style="List Bullet 2")

    # QC Table
    if rows:
        doc.add_heading("QC Board", level=1)

        table = doc.add_table(rows=1, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Type"
        hdr_cells[1].text = "Severity"
        hdr_cells[2].text = "Timestamp"
        hdr_cells[3].text = "Location"
        hdr_cells[4].text = "Snippet"
        hdr_cells[5].text = "Issue"
        hdr_cells[6].text = "Suggestion"

        for row in rows:
            try:
                cells = table.add_row().cells
                cells[0].text = _safe_str(row.get("Type"))
                cells[1].text = _safe_str(row.get("Severity"))
                cells[2].text = _safe_str(row.get("Timestamp"), "N/A")
                cells[3].text = _safe_str(row.get("Location"))
                cells[4].text = _safe_str(row.get("Snippet"))
                cells[5].text = _safe_str(row.get("Issue"))
                cells[6].text = _safe_str(row.get("Suggestion"))
            except Exception:
                continue  # Skip a malformed row rather than crashing the whole export

    # Transcript
    doc.add_heading("Transcript", level=1)
    doc.add_paragraph(transcript)

    # Save to bytes
    try:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as exc:
        # If saving fails (e.g. corrupt cell data), return a minimal error doc
        fallback = Document()
        fallback.add_heading("QC Report — Export Error", 0)
        fallback.add_paragraph(
            f"The full report could not be generated due to an error: {exc}. "
            "Please download the CSV or JSON export instead."
        )
        buf2 = BytesIO()
        fallback.save(buf2)
        buf2.seek(0)
        return buf2.getvalue()